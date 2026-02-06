from docx import Document
from docx.shared import Inches

doc = Document()
doc.add_heading('Sample Document', 0)

doc.add_heading('Introduction', 1)
doc.add_paragraph('This is a sample DOCX document with various formatting to test Markdown conversion.')

doc.add_heading('Features', 1)
p = doc.add_paragraph('This document includes:')
p.add_run(' bold text').bold = True
p.add_run(', ')
p.add_run('italic text').italic = True
p.add_run(', and ')
p.add_run('underlined text').underline = True

doc.add_heading('Lists', 2)
doc.add_paragraph('Unordered list:', style='List Bullet')
doc.add_paragraph('First item', style='List Bullet')
doc.add_paragraph('Second item', style='List Bullet')

doc.add_paragraph('Ordered list:', style='List Number')
doc.add_paragraph('First step', style='List Number')
doc.add_paragraph('Second step', style='List Number')

doc.add_heading('Table', 2)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Age'
hdr_cells[2].text = 'City'

row_cells = table.add_row().cells
row_cells[0].text = 'John'
row_cells[1].text = '30'
row_cells[2].text = 'New York'

row_cells = table.add_row().cells
row_cells[0].text = 'Jane'
row_cells[1].text = '25'
row_cells[2].text = 'London'

doc.add_heading('Conclusion', 1)
doc.add_paragraph('This concludes the sample document.')

doc.save('test.docx')
